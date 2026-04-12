"""Render cover letters to deterministic .docx format."""

from __future__ import annotations

from pathlib import Path  # noqa: TC003
from typing import TYPE_CHECKING, Any

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

if TYPE_CHECKING:
    from pipelines.job_agent.cover_letter.models import CoverLetterDocument

logger = structlog.get_logger(__name__)

_FONT_NAME = "Times New Roman"
_FONT_SIZE = Pt(11)


def render_cover_letter_docx(
    doc: CoverLetterDocument,
    output_path: Path,
    *,
    signature_name: str,
    sender_name: str,
    sender_location: str,
    sender_email: str,
    sender_phone: str,
) -> Path:
    """Render a polished business-letter .docx file."""
    document = Document()
    _setup_page(document)

    _add_line(document, sender_name, bold=True)
    if sender_location:
        _add_line(document, sender_location)
    if sender_email:
        _add_line(document, sender_email)
    if sender_phone:
        _add_line(document, sender_phone)
    _add_blank(document)

    _add_line(document, doc.salutation)
    _add_blank(document)

    _add_paragraph(document, doc.opening_paragraph)
    _add_blank(document)

    for paragraph in doc.body_paragraphs:
        _add_paragraph(document, paragraph)
        _add_blank(document)

    _add_paragraph(document, doc.closing_paragraph)
    _add_blank(document)

    _add_line(document, doc.signoff)
    _add_blank(document)
    _add_line(document, signature_name or sender_name)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    logger.info("cover_letter.rendered", path=str(output_path))

    preview_path = output_path.with_suffix(".md")
    _render_markdown_preview(
        doc,
        preview_path,
        sender_name=sender_name,
        sender_location=sender_location,
        sender_email=sender_email,
        sender_phone=sender_phone,
    )

    return output_path


def _setup_page(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def _add_line(document: Any, text: str, *, bold: bool = False) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE
    _zero_spacing(para)


def _add_paragraph(document: Any, text: str) -> None:
    para = document.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = _FONT_NAME
        run.font.size = _FONT_SIZE
    _zero_spacing(para)


def _add_blank(document: Any) -> None:
    para = document.add_paragraph("")
    _zero_spacing(para)


def _zero_spacing(para: Any) -> None:
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.15


# -- markdown preview ------------------------------------------------------


def _render_markdown_preview(
    doc: CoverLetterDocument,
    path: Path,
    *,
    sender_name: str,
    sender_location: str,
    sender_email: str,
    sender_phone: str,
) -> None:
    """Write a Markdown preview of the cover letter alongside the .docx."""
    lines: list[str] = []

    lines.append(f"**{sender_name}**")
    contact = " | ".join(p for p in [sender_location, sender_email, sender_phone] if p)
    if contact:
        lines.append(f"_{contact}_")
    lines.append("")
    lines.append("---")
    lines.append("")

    lines.append(doc.salutation)
    lines.append("")
    lines.append(doc.opening_paragraph)
    lines.append("")
    for para in doc.body_paragraphs:
        lines.append(para)
        lines.append("")
    lines.append(doc.closing_paragraph)
    lines.append("")
    lines.append(doc.signoff)
    lines.append("")
    lines.append(doc.signature_name)
    lines.append("")

    try:
        path.write_text("\n".join(lines), encoding="utf-8")
        logger.info("cover_letter_preview_rendered", path=str(path))
    except Exception:
        logger.warning("cover_letter_preview_failed", path=str(path), exc_info=True)
