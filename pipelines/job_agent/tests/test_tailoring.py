"""Tests for the resume tailoring pipeline.

Covers profile loading, plan application, .docx rendering, context
pruning, tag expansion, and the tailoring node with MockLLMClient
(no real API calls).

The job analysis pass is now a separate upstream node; tailoring
consumes pre-computed ``JobAnalysisResult`` from ``state.job_analyses``.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from core.testing import MockLLMClient
from pipelines.job_agent.models import ApplicationStatus, JobListing, JobSource, SearchCriteria
from pipelines.job_agent.models.resume_tailoring import (
    BulletOp,
    JobAnalysisResult,
    ResumeMasterProfile,
    ResumeTailoringPlan,
    SectionPlan,
    TailoredExperience,
    TailoredResumeDocument,
)
from pipelines.job_agent.nodes.tailoring import _expand_domain_tags
from pipelines.job_agent.resume.applier import apply_tailoring_plan
from pipelines.job_agent.resume.profile import format_profile_for_llm, load_master_profile
from pipelines.job_agent.resume.renderer import render_resume_docx
from pipelines.job_agent.state import JobAgentState, PipelinePhase

_FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ── fixtures ───────────────────────────────────────────────────────────


@pytest.fixture
def master_profile() -> ResumeMasterProfile:
    """Load the test fixture master profile."""
    return load_master_profile(_FIXTURES_DIR / "master_profile.yaml")


@pytest.fixture
def basic_plan() -> ResumeTailoringPlan:
    """A minimal valid tailoring plan referencing the test fixture."""
    return ResumeTailoringPlan(
        summary="Experienced engineer with ML and platform expertise.",
        experience_sections=[
            SectionPlan(section_id="exp_alpha", bullet_order=["alpha_platform", "alpha_ml"]),
            SectionPlan(section_id="exp_beta", bullet_order=["beta_pipeline"]),
        ],
        education_sections=[
            SectionPlan(section_id="edu_test", bullet_order=["edu_test_ml"]),
        ],
        bullet_ops=[
            BulletOp(bullet_id="alpha_platform", op="keep"),
            BulletOp(bullet_id="alpha_ml", op="keep"),
            BulletOp(bullet_id="beta_pipeline", op="shorten"),
            BulletOp(bullet_id="edu_test_ml", op="keep"),
        ],
        skills_to_highlight=["Python", "PyTorch", "Docker"],
    )


@pytest.fixture
def tailored_doc(
    master_profile: ResumeMasterProfile,
    basic_plan: ResumeTailoringPlan,
) -> TailoredResumeDocument:
    """A fully applied tailored document ready for rendering."""
    return apply_tailoring_plan(master_profile, basic_plan)


def _make_listing(
    *,
    description: str = "Build autonomous systems for defense...",
    dedup_key: str = "test_dedup_001",
) -> JobListing:
    return JobListing(
        title="Senior Technical Product Manager",
        company="Anduril Industries",
        location="Costa Mesa, CA",
        url="https://jobs.lever.co/anduril/abc123",
        source=JobSource.LEVER,
        description=description,
        salary_min=180_000,
        salary_max=250_000,
        remote=False,
        dedup_key=dedup_key,
    )


def _make_analysis() -> JobAnalysisResult:
    return JobAnalysisResult(
        themes=["autonomous systems", "defense technology", "product leadership"],
        seniority="senior",
        domain_tags=["defense", "tech", "product"],
        must_hit_keywords=["autonomous", "defense", "product management"],
        priority_requirements=["5+ years engineering", "defense background"],
        basic_qualifications=["BS in CS or equivalent"],
        preferred_qualifications=["Clearance preferred"],
        angles=["defense engineering to product", "technical depth"],
    )


def _mock_analysis_json() -> str:
    return _make_analysis().model_dump_json()


def _mock_plan_json() -> str:
    return json.dumps(
        {
            "summary": "Technical leader with defense and ML platform experience.",
            "experience_sections": [
                {"section_id": "exp_alpha", "bullet_order": ["alpha_platform", "alpha_cicd"]},
                {"section_id": "exp_beta", "bullet_order": ["beta_pipeline"]},
            ],
            "education_sections": [
                {"section_id": "edu_test", "bullet_order": ["edu_test_ml"]},
            ],
            "bullet_ops": [
                {"bullet_id": "alpha_platform", "op": "keep"},
                {"bullet_id": "alpha_cicd", "op": "shorten"},
                {"bullet_id": "beta_pipeline", "op": "keep"},
                {"bullet_id": "edu_test_ml", "op": "keep"},
            ],
            "skills_to_highlight": ["Python", "Docker", "AWS"],
        }
    )


# ── Profile Loading ────────────────────────────────────────────────────


class TestMasterProfileLoading:
    """Tests for YAML profile loading and validation."""

    def test_load_valid_profile(self, master_profile: ResumeMasterProfile) -> None:
        assert master_profile.name == "Test Candidate"
        assert master_profile.schema_version == 1
        assert len(master_profile.experience) == 2
        assert len(master_profile.education) == 1

    def test_all_bullet_ids(self, master_profile: ResumeMasterProfile) -> None:
        ids = master_profile.all_bullet_ids()
        assert "alpha_platform" in ids
        assert "beta_fraud" in ids
        assert "edu_test_ml" in ids
        assert len(ids) == 8  # 4 alpha + 2 beta + 2 edu

    def test_get_bullet(self, master_profile: ResumeMasterProfile) -> None:
        b = master_profile.get_bullet("alpha_platform")
        assert b is not None
        assert "microservices" in b.text
        assert "technical" in b.tags

    def test_get_bullet_returns_none_for_unknown(self, master_profile: ResumeMasterProfile) -> None:
        assert master_profile.get_bullet("nonexistent") is None

    def test_get_experience(self, master_profile: ResumeMasterProfile) -> None:
        exp = master_profile.get_experience("exp_alpha")
        assert exp is not None
        assert exp.company == "Alpha Corp"

    def test_get_education(self, master_profile: ResumeMasterProfile) -> None:
        edu = master_profile.get_education("edu_test")
        assert edu is not None
        assert edu.school == "Test University"

    def test_file_not_found(self) -> None:
        with pytest.raises(FileNotFoundError):
            load_master_profile(Path("/nonexistent/profile.yaml"))

    def test_format_for_llm_contains_ids(self, master_profile: ResumeMasterProfile) -> None:
        text = format_profile_for_llm(master_profile)
        assert "[alpha_platform]" in text
        assert "[exp_alpha]" in text
        assert "[edu_test]" in text
        assert "EXPERIENCE:" in text
        assert "EDUCATION:" in text
        assert "SKILLS:" in text

    def test_format_with_tag_filter_includes_matching(
        self, master_profile: ResumeMasterProfile
    ) -> None:
        text = format_profile_for_llm(master_profile, relevant_tags={"ml", "technical"})
        assert "[alpha_platform]" in text
        assert "[alpha_ml]" in text
        assert "[edu_test_ml]" in text

    def test_format_with_tag_filter_excludes_non_matching(
        self, master_profile: ResumeMasterProfile
    ) -> None:
        text = format_profile_for_llm(master_profile, relevant_tags={"finance"})
        assert "[beta_fraud]" in text
        assert "[alpha_platform]" not in text
        assert "[alpha_ml]" not in text

    def test_format_with_tag_filter_omits_empty_sections(
        self, master_profile: ResumeMasterProfile
    ) -> None:
        text = format_profile_for_llm(master_profile, relevant_tags={"finance"})
        assert "[exp_alpha]" not in text
        assert "[exp_beta]" in text

    def test_format_with_none_tags_includes_all(self, master_profile: ResumeMasterProfile) -> None:
        full = format_profile_for_llm(master_profile)
        also_full = format_profile_for_llm(master_profile, relevant_tags=None)
        assert full == also_full

    def test_filtered_profile_is_shorter(self, master_profile: ResumeMasterProfile) -> None:
        full = format_profile_for_llm(master_profile)
        filtered = format_profile_for_llm(master_profile, relevant_tags={"finance"})
        assert len(filtered) < len(full)


# ── Tag Expansion ─────────────────────────────────────────────────────


class TestTagExpansion:
    """Tests for domain tag expansion to profile tag vocabulary."""

    def test_direct_tags_preserved(self) -> None:
        result = _expand_domain_tags(["defense", "ml"])
        assert "defense" in result
        assert "ml" in result

    def test_synonyms_expanded(self) -> None:
        result = _expand_domain_tags(["military"])
        assert "defense" in result
        assert "naval" in result

    def test_always_relevant_tags_included(self) -> None:
        result = _expand_domain_tags(["finance"])
        assert "leadership" in result
        assert "technical" in result
        assert "management" in result
        assert "software" in result

    def test_case_insensitive(self) -> None:
        result = _expand_domain_tags(["Defense", "ML"])
        assert "defense" in result
        assert "ml" in result


# ── Applier ────────────────────────────────────────────────────────────


class TestApplier:
    """Tests for the plan application logic (pure, no LLM)."""

    def test_apply_basic_plan(
        self,
        master_profile: ResumeMasterProfile,
        basic_plan: ResumeTailoringPlan,
    ) -> None:
        doc = apply_tailoring_plan(master_profile, basic_plan)
        assert doc.name == "Test Candidate"
        assert doc.summary == "Experienced engineer with ML and platform expertise."
        assert len(doc.experience) == 2
        assert len(doc.education) == 1
        assert doc.skills_highlight == ["Python", "PyTorch", "Docker"]

    def test_bullet_ordering(
        self,
        master_profile: ResumeMasterProfile,
        basic_plan: ResumeTailoringPlan,
    ) -> None:
        doc = apply_tailoring_plan(master_profile, basic_plan)
        alpha_bullets = doc.experience[0].bullets
        assert alpha_bullets[0].id == "alpha_platform"
        assert alpha_bullets[1].id == "alpha_ml"

    def test_shorten_op_uses_variant(
        self,
        master_profile: ResumeMasterProfile,
        basic_plan: ResumeTailoringPlan,
    ) -> None:
        doc = apply_tailoring_plan(master_profile, basic_plan)
        beta_bullet = doc.experience[1].bullets[0]
        assert beta_bullet.id == "beta_pipeline"
        assert beta_bullet.text == "Real-time pipeline (500K events/sec)"

    def test_rewrite_op(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Test summary.",
            experience_sections=[
                SectionPlan(section_id="exp_alpha", bullet_order=["alpha_platform"]),
            ],
            education_sections=[],
            bullet_ops=[
                BulletOp(
                    bullet_id="alpha_platform",
                    op="rewrite",
                    rewrite_text="Custom rewritten bullet text",
                ),
            ],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert doc.experience[0].bullets[0].text == "Custom rewritten bullet text"

    def test_unknown_section_skipped(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Test.",
            experience_sections=[
                SectionPlan(section_id="nonexistent_section", bullet_order=["alpha_platform"]),
            ],
            education_sections=[],
            bullet_ops=[],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert len(doc.experience) == 0

    def test_unknown_bullet_skipped(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Test.",
            experience_sections=[
                SectionPlan(section_id="exp_alpha", bullet_order=["nonexistent_bullet"]),
            ],
            education_sections=[],
            bullet_ops=[],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert len(doc.experience[0].bullets) == 0

    def test_shorten_without_variant_falls_back(self, master_profile: ResumeMasterProfile) -> None:
        """Shorten op on a bullet with no short variant -> keeps original text."""
        plan = ResumeTailoringPlan(
            summary="Test.",
            experience_sections=[
                SectionPlan(section_id="exp_alpha", bullet_order=["alpha_mentoring"]),
            ],
            education_sections=[],
            bullet_ops=[BulletOp(bullet_id="alpha_mentoring", op="shorten")],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        original = master_profile.get_bullet("alpha_mentoring")
        assert original is not None
        assert doc.experience[0].bullets[0].text == original.text

    def test_normalizes_inline_dash_in_prose(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Summary with em dash \u2014 still fine.",
            experience_sections=[
                SectionPlan(section_id="exp_alpha", bullet_order=["alpha_platform"]),
            ],
            education_sections=[],
            bullet_ops=[
                BulletOp(
                    bullet_id="alpha_platform",
                    op="rewrite",
                    rewrite_text="Rewritten bullet with em dash \u2014 remove it",
                ),
            ],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert "\u2014" not in doc.summary
        assert "\u2014" not in doc.experience[0].bullets[0].text
        assert "; " in doc.summary
        assert "; " in doc.experience[0].bullets[0].text

    def test_preserves_dash_like_metadata(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Normal summary text.",
            experience_sections=[
                SectionPlan(section_id="exp_alpha", bullet_order=["alpha_platform"]),
            ],
            education_sections=[],
            bullet_ops=[BulletOp(bullet_id="alpha_platform", op="keep")],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert doc.experience[0].title == "Senior Engineer"
        assert doc.experience[0].dates == "2022-2024"

    def test_location_and_subtitle_passthrough(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Test.",
            experience_sections=[
                SectionPlan(section_id="exp_alpha", bullet_order=["alpha_platform"]),
            ],
            education_sections=[
                SectionPlan(section_id="edu_test", bullet_order=["edu_test_ml"]),
            ],
            bullet_ops=[
                BulletOp(bullet_id="alpha_platform", op="keep"),
                BulletOp(bullet_id="edu_test_ml", op="keep"),
            ],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert doc.experience[0].location == "New York, NY"
        assert doc.experience[0].subtitle == "Enterprise SaaS Platform"
        assert doc.education[0].location == "Boston, MA"

    def test_additional_info_includes_clearance(self, master_profile: ResumeMasterProfile) -> None:
        plan = ResumeTailoringPlan(
            summary="Test.",
            experience_sections=[],
            education_sections=[],
            bullet_ops=[],
            skills_to_highlight=[],
        )
        doc = apply_tailoring_plan(master_profile, plan)
        assert any("Test Clearance" in item for item in doc.additional_info)


# ── Renderer ───────────────────────────────────────────────────────────


class TestRenderer:
    """Tests for .docx rendering matching the Kokomoor template format."""

    def test_render_creates_file(
        self, tailored_doc: TailoredResumeDocument, tmp_path: Path
    ) -> None:
        out = tmp_path / "resume.docx"
        result = render_resume_docx(tailored_doc, out)
        assert result.exists()
        assert result.stat().st_size > 0

    def test_render_creates_parent_dirs(
        self, tailored_doc: TailoredResumeDocument, tmp_path: Path
    ) -> None:
        out = tmp_path / "nested" / "dir" / "resume.docx"
        result = render_resume_docx(tailored_doc, out)
        assert result.exists()

    def test_rendered_doc_has_content(
        self, tailored_doc: TailoredResumeDocument, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "resume.docx"
        render_resume_docx(tailored_doc, out)
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Test Candidate" in t for t in texts)

    def test_education_before_experience(
        self, tailored_doc: TailoredResumeDocument, tmp_path: Path
    ) -> None:
        from docx import Document

        out = tmp_path / "resume.docx"
        render_resume_docx(tailored_doc, out)
        doc = Document(str(out))
        texts = [p.text.lower() for p in doc.paragraphs if p.text.strip()]
        edu_idx = next(i for i, t in enumerate(texts) if "education" in t)
        exp_idx = next(i for i, t in enumerate(texts) if "experience" in t)
        assert edu_idx < exp_idx

    def test_location_and_subtitle_rendered(self, tmp_path: Path) -> None:
        from docx import Document

        from pipelines.job_agent.models.resume_tailoring import TailoredBullet

        doc = TailoredResumeDocument(
            name="Test",
            location="Boston, MA",
            email="t@t.com",
            phone="555",
            linkedin="",
            github="",
            clearance="",
            summary="",
            experience=[
                TailoredExperience(
                    company="Acme",
                    title="Lead",
                    dates="2024",
                    location="NYC",
                    subtitle="Defense Contractor",
                    bullets=[TailoredBullet(id="b1", text="Did work")],
                )
            ],
            education=[],
            skills_highlight=[],
        )
        out = tmp_path / "resume.docx"
        render_resume_docx(doc, out)
        rendered = Document(str(out))
        texts = [p.text for p in rendered.paragraphs]
        assert any("NYC" in t for t in texts)
        assert any("Defense Contractor" in t for t in texts)


# ── Tailoring Node ─────────────────────────────────────────────────────


class TestTailoringNode:
    """Integration tests for the tailoring node with MockLLMClient.

    The tailoring node now expects ``state.job_analyses`` to be
    pre-populated by the upstream job-analysis node. It only runs
    the plan pass (1 LLM call per listing).
    """

    @pytest.mark.asyncio
    async def test_tailors_qualified_listings(self, tmp_path: Path) -> None:
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        mock_client = MockLLMClient(responses=[_mock_plan_json()])
        listing = _make_listing()
        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[listing],
            job_analyses={listing.dedup_key: _make_analysis()},
            run_id="test-run",
        )

        _patch_settings(tmp_path)
        result = await tailoring_node(state, llm_client=mock_client)

        assert result.phase == PipelinePhase.TAILORING
        assert result.tailored_listings is result.qualified_listings
        assert len(result.errors) == 0

        assert listing.tailored_resume_path is not None
        assert Path(listing.tailored_resume_path).exists()
        assert listing.status == ApplicationStatus.PENDING_REVIEW

    @pytest.mark.asyncio
    async def test_skips_dry_run(self) -> None:
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[_make_listing()],
            run_id="test-dry",
            dry_run=True,
        )
        result = await tailoring_node(state)
        assert result.tailored_listings is result.qualified_listings
        assert result.qualified_listings[0].tailored_resume_path is None

    @pytest.mark.asyncio
    async def test_skips_empty_listings(self) -> None:
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[],
            run_id="test-empty",
        )
        result = await tailoring_node(state)
        assert result.tailored_listings == []

    @pytest.mark.asyncio
    async def test_errors_on_missing_analysis(self, tmp_path: Path) -> None:
        """Listing without a matching analysis produces an error, not a crash."""
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        mock_client = MockLLMClient(responses=[])
        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[_make_listing()],
            job_analyses={},
            run_id="test-missing-analysis",
        )
        _patch_settings(tmp_path)
        result = await tailoring_node(state, llm_client=mock_client)

        assert len(result.errors) == 1
        assert "no job analysis" in result.errors[0]["message"].lower()
        assert result.qualified_listings[0].status == ApplicationStatus.ERRORED

    @pytest.mark.asyncio
    async def test_sets_errored_status_when_tailoring_fails(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        def _boom(*args: object, **kwargs: object) -> None:
            raise RuntimeError("render failed")

        monkeypatch.setattr("pipelines.job_agent.nodes.tailoring.render_resume_docx", _boom)
        mock_client = MockLLMClient(responses=[_mock_plan_json()])
        listing = _make_listing()
        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[listing],
            job_analyses={listing.dedup_key: _make_analysis()},
            run_id="test-tailoring-error",
        )
        _patch_settings(tmp_path)
        result = await tailoring_node(state, llm_client=mock_client)
        assert result.errors
        assert listing.status == ApplicationStatus.ERRORED

    @pytest.mark.asyncio
    async def test_one_llm_call_per_listing(self, tmp_path: Path) -> None:
        """Only the plan pass calls the LLM (analysis is pre-computed)."""
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        mock_client = MockLLMClient(responses=[_mock_plan_json()])
        listing = _make_listing()
        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[listing],
            job_analyses={listing.dedup_key: _make_analysis()},
            run_id="test-calls",
        )
        _patch_settings(tmp_path)
        await tailoring_node(state, llm_client=mock_client)

        assert len(mock_client.calls) == 1

    @pytest.mark.asyncio
    async def test_plan_model_from_config(self, tmp_path: Path) -> None:
        from pipelines.job_agent.nodes.tailoring import tailoring_node

        mock_client = MockLLMClient(responses=[_mock_plan_json()])
        listing = _make_listing()
        state = JobAgentState(
            search_criteria=SearchCriteria(),
            qualified_listings=[listing],
            job_analyses={listing.dedup_key: _make_analysis()},
            run_id="test-plan-model",
        )
        _patch_settings(tmp_path)
        await tailoring_node(state, llm_client=mock_client)

        plan_call = mock_client.calls[0]
        assert plan_call[1]["model"] is None  # uses client default


@pytest.mark.asyncio
async def test_analysis_then_tailoring_preserves_state_mutation_pattern(tmp_path: Path) -> None:
    """Regression: analysis + tailoring still mutates listings in place and writes paths."""
    from pipelines.job_agent.nodes.job_analysis import job_analysis_node
    from pipelines.job_agent.nodes.tailoring import tailoring_node

    listing = _make_listing(description="A" * 5000)
    state = JobAgentState(
        search_criteria=SearchCriteria(),
        qualified_listings=[listing],
        run_id="test-regression-flow",
    )
    _patch_settings(tmp_path)
    mock_client = MockLLMClient(responses=[_mock_analysis_json(), _mock_plan_json()])

    after_analysis = await job_analysis_node(state, llm_client=mock_client)
    result = await tailoring_node(after_analysis, llm_client=mock_client)

    assert len(mock_client.calls) == 2
    assert result.tailored_listings is result.qualified_listings
    assert result.qualified_listings[0] is listing
    assert listing.tailored_resume_path is not None
    assert Path(listing.tailored_resume_path).exists()
    assert listing.status == ApplicationStatus.PENDING_REVIEW


# ── test helpers ───────────────────────────────────────────────────────


def _patch_settings(
    tmp_path: Path,
    *,
    plan_model: str = "",
    plan_max_tokens: int = 2048,
) -> None:
    """Point resume settings at the test fixture profile and tmp output dir."""
    import os

    from core.config import get_settings

    get_settings.cache_clear()
    os.environ["KP_RESUME_MASTER_PROFILE_PATH"] = str(_FIXTURES_DIR / "master_profile.yaml")
    os.environ["KP_RESUME_OUTPUT_DIR"] = str(tmp_path / "output")
    os.environ["KP_RESUME_PLAN_MODEL"] = plan_model
    os.environ["KP_RESUME_PLAN_MAX_TOKENS"] = str(plan_max_tokens)
    os.environ["KP_JOB_ANALYSIS_MODEL"] = "claude-haiku-4-5-20251001"
    os.environ["KP_JOB_ANALYSIS_ENABLE_CACHE"] = "true"
    get_settings.cache_clear()
