"""Tests for cover-letter tailoring specialization."""

from __future__ import annotations

import json
import os
from pathlib import Path

import pytest
from docx import Document

from core.config import get_settings
from core.testing import MockLLMClient
from pipelines.job_agent.cover_letter.profile import (
    format_cover_letter_inventory,
    load_cover_letter_style_guide,
)
from pipelines.job_agent.cover_letter.renderer import render_cover_letter_docx
from pipelines.job_agent.cover_letter.validation import validate_cover_letter_plan
from pipelines.job_agent.models import ApplicationStatus, JobListing, JobSource, SearchCriteria
from pipelines.job_agent.models.resume_tailoring import JobAnalysisResult
from pipelines.job_agent.resume.profile import load_master_profile
from pipelines.job_agent.state import JobAgentState, PipelinePhase

_FIXTURES_DIR = Path(__file__).parent / "fixtures"


def _make_listing(*, dedup_key: str = "cover_key_001") -> JobListing:
    return JobListing(
        title="Senior Product Manager, Autonomy",
        company="Anduril Industries",
        location="Costa Mesa, CA",
        url="https://jobs.example.com/anduril/autonomy",
        source=JobSource.LEVER,
        description="Lead autonomous systems product strategy and cross-functional delivery.",
        dedup_key=dedup_key,
    )


def _make_analysis() -> JobAnalysisResult:
    return JobAnalysisResult(
        themes=["autonomous systems", "defense", "product leadership"],
        seniority="senior",
        domain_tags=["defense", "ai", "product"],
        must_hit_keywords=["autonomy", "systems engineering", "mission"],
        priority_requirements=["leadership", "technical depth", "cross-functional execution"],
        basic_qualifications=["BS degree"],
        preferred_qualifications=["Defense experience"],
        angles=["operator empathy", "technical credibility"],
    )


def _mock_cover_letter_plan_json() -> str:
    return json.dumps(
        {
            "salutation": "Dear Hiring Manager,",
            "opening_paragraph": (
                "The Senior Product Manager, Autonomy role at Anduril Industries sits at the "
                "intersection of what I have spent my career building: reliable technical systems "
                "in high-stakes environments where execution discipline matters."
            ),
            "body_paragraphs": [
                (
                    "At Alpha Corp, I led development of a microservices platform serving 10M+ requests "
                    "per day across 15 services, and improved deployment velocity by 70% through CI/CD "
                    "automation. That experience maps directly to shipping reliable autonomy capabilities."
                ),
                (
                    "I also built an ML-powered anomaly detection system that reduced incident response "
                    "time by 40%, and previously built a real-time data pipeline processing 500K events per "
                    "second. Those outcomes reflect the systems mindset needed for complex mission software."
                ),
            ],
            "closing_paragraph": (
                "Anduril Industries stands out because it applies advanced engineering to urgent national "
                "security needs. I would value the chance to help translate operator needs into deployable "
                "autonomy products."
            ),
            "signoff": "Sincerely,",
            "signature_name": "Test Candidate",
            "company_motivation": (
                "Anduril Industries applies advanced engineering to urgent national security "
                "challenges, building autonomous systems that protect service members in the field."
            ),
            "job_requirements_addressed": [
                "technical leadership",
                "reliable delivery",
                "ml systems",
            ],
            "selected_experience_ids": ["exp_alpha", "exp_beta"],
            "selected_bullet_ids": ["alpha_platform", "alpha_cicd", "alpha_ml", "beta_pipeline"],
            "selected_education_ids": ["edu_test"],
            "requirement_evidence": [
                {
                    "requirement": "Lead technical teams and ship reliable systems",
                    "supporting_bullet_ids": ["alpha_platform", "alpha_cicd"],
                },
                {
                    "requirement": "Apply machine learning in production",
                    "supporting_bullet_ids": ["alpha_ml", "beta_pipeline"],
                },
            ],
            "tone_version": "confident_direct",
        }
    )


def _patch_settings(tmp_path: Path) -> None:
    get_settings.cache_clear()
    os.environ["KP_RESUME_MASTER_PROFILE_PATH"] = str(_FIXTURES_DIR / "master_profile.yaml")
    os.environ["KP_COVER_LETTER_OUTPUT_DIR"] = str(tmp_path / "cover_letters")
    os.environ["KP_COVER_LETTER_MAX_TOKENS"] = "2200"
    os.environ["KP_COVER_LETTER_MAX_INPUT_CHARS"] = "12000"
    os.environ["KP_COVER_LETTER_MODEL"] = "claude-sonnet-4-20250514"
    os.environ["KP_COVER_LETTER_STYLE_GUIDE_PATH"] = str(
        Path("pipelines/job_agent/context/cover_letter_style.md")
    )
    get_settings.cache_clear()


def test_style_guide_fallback_when_missing(tmp_path: Path) -> None:
    content = load_cover_letter_style_guide(tmp_path / "missing.md")
    assert "Cover Letter Style Guide" in content


def test_style_guide_loads_from_path(tmp_path: Path) -> None:
    style = tmp_path / "guide.md"
    style.write_text("# My Style\nUse concrete examples.", encoding="utf-8")
    content = load_cover_letter_style_guide(style)
    assert "My Style" in content


def test_inventory_includes_cover_letter_preferences() -> None:
    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    text = format_cover_letter_inventory(profile)
    assert "COVER LETTER PREFERENCES:" in text
    assert "Preferred signoff" in text


def test_inventory_graceful_without_cover_letter_section(tmp_path: Path) -> None:
    profile_text = (_FIXTURES_DIR / "master_profile.yaml").read_text(encoding="utf-8")
    without_cover = profile_text.split("\ncover_letter:", maxsplit=1)[0]
    no_cover_path = tmp_path / "master_profile_no_cover.yaml"
    no_cover_path.write_text(without_cover, encoding="utf-8")
    profile = load_master_profile(no_cover_path)
    text = format_cover_letter_inventory(profile)
    assert "COVER LETTER PREFERENCES" not in text


def test_inventory_tag_filtering() -> None:
    """Context pruning omits bullets whose tags don't overlap with relevant_tags."""
    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    full_text = format_cover_letter_inventory(profile)
    pruned_text = format_cover_letter_inventory(profile, relevant_tags={"ml"})
    assert len(pruned_text) < len(full_text)
    assert "alpha_ml" in pruned_text
    assert "alpha_cicd" not in pruned_text


def test_validation_happy_path() -> None:
    """Full validation succeeds on a well-formed plan with no warnings."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    result = validate_cover_letter_plan(
        plan=plan,
        profile=profile,
        expected_company="Anduril Industries",
    )
    assert result.document.salutation == "Dear Hiring Manager,"
    assert result.document.signoff == "Sincerely,"
    assert len(result.document.body_paragraphs) == 2
    assert len(result.plan.requirement_evidence) == 2
    assert result.plan.requirement_evidence[0].requirement
    assert result.plan.tone_version == "confident_direct"


def test_validation_rejects_unknown_ids() -> None:
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.selected_bullet_ids.append("missing_id")

    with pytest.raises(ValueError, match="unknown profile IDs"):
        validate_cover_letter_plan(
            plan=plan, profile=profile, expected_company="Anduril Industries"
        )


def test_validation_normalizes_dashes_and_signoff() -> None:
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.body_paragraphs[0] = "This paragraph has em dash \u2014 that must be cleaned."
    plan.signoff = "Regards"
    result = validate_cover_letter_plan(
        plan=plan,
        profile=profile,
        expected_company="Anduril Industries",
    )
    assert "\u2014" not in result.document.body_paragraphs[0]
    assert result.document.signoff.endswith(",")


def test_validation_rejects_placeholders() -> None:
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.opening_paragraph = "The role at [Company] is compelling because of its mission focus."

    with pytest.raises(ValueError, match="placeholder"):
        validate_cover_letter_plan(
            plan=plan, profile=profile, expected_company="Anduril Industries"
        )


def test_validation_rejects_core_banned_phrase() -> None:
    """Core banned phrases cause hard failure regardless of preferences."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.body_paragraphs[0] = (
        "I am passionate about building reliable systems. At Alpha Corp, I led development "
        "of a microservices platform serving 10M+ requests per day across 15 services."
    )

    with pytest.raises(ValueError, match="banned phrase"):
        validate_cover_letter_plan(
            plan=plan, profile=profile, expected_company="Anduril Industries"
        )


def test_validation_rejects_profile_banned_phrase() -> None:
    """Profile-level banned phrases also cause hard failure."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    assert profile.cover_letter is not None
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.body_paragraphs[0] = (
        "I want to leverage my background to build microservices platform serving 10M+ "
        "requests per day and improve deployment velocity by 70% through CI/CD automation."
    )

    with pytest.raises(ValueError, match="profile-banned phrase"):
        validate_cover_letter_plan(
            plan=plan,
            profile=profile,
            expected_company="Anduril Industries",
            preferences=profile.cover_letter,
        )


def test_validation_applies_preferred_signoff() -> None:
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    assert profile.cover_letter is not None
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.signoff = "Best"

    result = validate_cover_letter_plan(
        plan=plan,
        profile=profile,
        expected_company="Anduril Industries",
        preferences=profile.cover_letter,
    )
    assert result.plan.signoff == "Sincerely,"


def test_validation_warns_on_missing_company_in_body() -> None:
    """Company missing from letter body produces a warning."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.opening_paragraph = (
        "I am applying for the Senior Product Manager, Autonomy role. "
        "I have led technical programs where mission urgency mattered."
    )
    plan.closing_paragraph = (
        "I would value the chance to help translate operator needs into "
        "deployable autonomy products at this company."
    )

    result = validate_cover_letter_plan(
        plan=plan,
        profile=profile,
        expected_company="Anduril Industries",
    )
    assert any("body does not mention" in w for w in result.warnings)


def test_renderer_layout_order(tmp_path: Path) -> None:
    from pipelines.job_agent.cover_letter.models import CoverLetterDocument

    doc = CoverLetterDocument(
        salutation="Dear Hiring Manager,",
        opening_paragraph="Opening paragraph.",
        body_paragraphs=["Body paragraph one.", "Body paragraph two."],
        closing_paragraph="Closing paragraph.",
        signoff="Sincerely,",
        signature_name="Test Candidate",
    )
    out = tmp_path / "letter.docx"
    render_cover_letter_docx(
        doc,
        out,
        signature_name="Test Candidate",
        sender_name="Test Candidate",
        sender_location="Boston, MA",
        sender_email="test@example.com",
        sender_phone="555-0100",
    )
    assert out.exists()

    rendered = Document(str(out))
    texts = [p.text for p in rendered.paragraphs if p.text.strip()]
    assert texts.index("Dear Hiring Manager,") < texts.index("Opening paragraph.")
    assert texts[-2] == "Sincerely,"
    assert texts[-1] == "Test Candidate"


@pytest.mark.asyncio
async def test_cover_letter_node_integration(tmp_path: Path) -> None:
    from pipelines.job_agent.nodes.cover_letter_tailoring import cover_letter_tailoring_node

    _patch_settings(tmp_path)
    listing = _make_listing()
    state = JobAgentState(
        search_criteria=SearchCriteria(),
        qualified_listings=[listing],
        job_analyses={listing.dedup_key: _make_analysis()},
        run_id="test-cover-letter",
    )
    mock_client = MockLLMClient(responses=[_mock_cover_letter_plan_json()])

    result = await cover_letter_tailoring_node(state, llm_client=mock_client)

    assert result.phase == PipelinePhase.COVER_LETTER_TAILORING
    assert listing.tailored_cover_letter_path is not None
    assert Path(listing.tailored_cover_letter_path).exists()
    assert listing.status == ApplicationStatus.PENDING_REVIEW
    assert len(mock_client.calls) == 1


@pytest.mark.asyncio
async def test_cover_letter_prompt_respects_input_cap(tmp_path: Path) -> None:
    from pipelines.job_agent.nodes.cover_letter_tailoring import cover_letter_tailoring_node

    _patch_settings(tmp_path)
    os.environ["KP_COVER_LETTER_MAX_INPUT_CHARS"] = "2000"
    get_settings.cache_clear()

    listing = _make_listing(dedup_key="cap_test_001")
    listing.description = "A" * 5000
    state = JobAgentState(
        search_criteria=SearchCriteria(),
        qualified_listings=[listing],
        job_analyses={listing.dedup_key: _make_analysis()},
        run_id="test-cover-letter-cap",
    )
    mock_client = MockLLMClient(responses=[_mock_cover_letter_plan_json()])
    await cover_letter_tailoring_node(state, llm_client=mock_client)

    prompt = mock_client.calls[0][0]
    assert "A" * 2001 not in prompt


@pytest.mark.asyncio
async def test_cover_letter_node_missing_analysis_errors(tmp_path: Path) -> None:
    from pipelines.job_agent.nodes.cover_letter_tailoring import cover_letter_tailoring_node

    _patch_settings(tmp_path)
    listing = _make_listing()
    state = JobAgentState(
        search_criteria=SearchCriteria(),
        qualified_listings=[listing],
        job_analyses={},
        run_id="test-cover-letter-missing-analysis",
    )
    result = await cover_letter_tailoring_node(state, llm_client=MockLLMClient(responses=[]))

    assert result.errors
    assert listing.status == ApplicationStatus.ERRORED


@pytest.mark.asyncio
async def test_cover_letter_shape_snapshot(tmp_path: Path) -> None:
    """Regression test for structured plan shape stability."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    _patch_settings(tmp_path)
    parsed = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    assert parsed.model_dump() == {
        "salutation": "Dear Hiring Manager,",
        "opening_paragraph": (
            "The Senior Product Manager, Autonomy role at Anduril Industries sits at the "
            "intersection of what I have spent my career building: reliable technical systems "
            "in high-stakes environments where execution discipline matters."
        ),
        "body_paragraphs": [
            (
                "At Alpha Corp, I led development of a microservices platform serving 10M+ requests "
                "per day across 15 services, and improved deployment velocity by 70% through CI/CD "
                "automation. That experience maps directly to shipping reliable autonomy capabilities."
            ),
            (
                "I also built an ML-powered anomaly detection system that reduced incident response "
                "time by 40%, and previously built a real-time data pipeline processing 500K events per "
                "second. Those outcomes reflect the systems mindset needed for complex mission software."
            ),
        ],
        "closing_paragraph": (
            "Anduril Industries stands out because it applies advanced engineering to urgent national "
            "security needs. I would value the chance to help translate operator needs into deployable "
            "autonomy products."
        ),
        "signoff": "Sincerely,",
        "signature_name": "Test Candidate",
        "company_motivation": (
            "Anduril Industries applies advanced engineering to urgent national security "
            "challenges, building autonomous systems that protect service members in the field."
        ),
        "job_requirements_addressed": ["technical leadership", "reliable delivery", "ml systems"],
        "selected_experience_ids": ["exp_alpha", "exp_beta"],
        "selected_bullet_ids": ["alpha_platform", "alpha_cicd", "alpha_ml", "beta_pipeline"],
        "selected_education_ids": ["edu_test"],
        "requirement_evidence": [
            {
                "requirement": "Lead technical teams and ship reliable systems",
                "supporting_bullet_ids": ["alpha_platform", "alpha_cicd"],
            },
            {
                "requirement": "Apply machine learning in production",
                "supporting_bullet_ids": ["alpha_ml", "beta_pipeline"],
            },
        ],
        "tone_version": "confident_direct",
    }


def test_validation_rejects_generic_opener() -> None:
    """Letters starting with stock openers are rejected."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.opening_paragraph = (
        "I am writing to express my interest in the Senior Product Manager role "
        "at Anduril Industries. I have led microservices platform development."
    )

    with pytest.raises(ValueError, match="generic opener"):
        validate_cover_letter_plan(
            plan=plan, profile=profile, expected_company="Anduril Industries"
        )


def test_validation_rejects_shallow_company_motivation() -> None:
    """company_motivation with too few words is rejected."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.company_motivation = "I like Anduril."

    with pytest.raises(ValueError, match="company_motivation must contain"):
        validate_cover_letter_plan(
            plan=plan, profile=profile, expected_company="Anduril Industries"
        )


def test_validation_rejects_ungrounded_evidence() -> None:
    """Body that cites bullet IDs but contains no specific terms from them is rejected."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.body_paragraphs = [
        "I have strong experience in building systems and leading teams to success.",
        "My work has consistently delivered value across multiple organizations.",
    ]

    with pytest.raises(ValueError, match="lacks specific evidence"):
        validate_cover_letter_plan(
            plan=plan, profile=profile, expected_company="Anduril Industries"
        )


def test_validation_warns_on_ai_tell_density() -> None:
    """Multiple AI-tell words generate a warning."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan
    from pipelines.job_agent.cover_letter.validation import AI_TELL_WORDS

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    ai_words = list(AI_TELL_WORDS[:4])
    plan.body_paragraphs[0] = (
        f"I want to {ai_words[0]} into the {ai_words[1]} of microservices platform "
        f"serving 10M+ requests, a {ai_words[2]} achievement. This {ai_words[3]} "
        "outcome improved deployment velocity by 70% through CI/CD automation."
    )

    result = validate_cover_letter_plan(
        plan=plan, profile=profile, expected_company="Anduril Industries"
    )
    assert any("AI-tell words" in w for w in result.warnings)


def test_validation_warns_on_motivation_body_mismatch() -> None:
    """company_motivation substance not reflected in body produces warning."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())
    plan.company_motivation = (
        "Anduril Industries is revolutionizing counter-drone warfare and maritime "
        "surveillance through lattice-enabled autonomous platforms deployed globally."
    )

    result = validate_cover_letter_plan(
        plan=plan, profile=profile, expected_company="Anduril Industries"
    )
    assert any("company_motivation reasoning" in w for w in result.warnings)


def test_tone_version_rejects_invalid_value() -> None:
    """Invalid tone_version value fails Pydantic validation."""
    from pydantic import ValidationError

    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    raw = json.loads(_mock_cover_letter_plan_json())
    raw["tone_version"] = "casual_bro_v1"

    with pytest.raises(ValidationError):
        CoverLetterPlan.model_validate(raw)


def test_tone_version_accepts_all_valid_values() -> None:
    """All defined tone versions parse successfully."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    for tone in ("confident_direct", "professional_narrative", "technical_precise"):
        raw = json.loads(_mock_cover_letter_plan_json())
        raw["tone_version"] = tone
        parsed = CoverLetterPlan.model_validate(raw)
        assert parsed.tone_version == tone


def test_evidence_grounding_passes_with_specific_terms() -> None:
    """Body containing specific terms from cited bullets passes grounding check."""
    from pipelines.job_agent.cover_letter.models import CoverLetterPlan

    profile = load_master_profile(_FIXTURES_DIR / "master_profile.yaml")
    plan = CoverLetterPlan.model_validate_json(_mock_cover_letter_plan_json())

    result = validate_cover_letter_plan(
        plan=plan, profile=profile, expected_company="Anduril Industries"
    )
    assert not any("lacks specific evidence" in w for w in result.warnings)
