"""Render a ``TailoredResumeDocument`` to a .docx file.

Produces a resume matching the Kokomoor template format:
- Times New Roman 11.5pt throughout
- Tight 0.5"/0.65" margins, 10pt minimum line spacing
- Section headers: bold, ALL CAPS, black bottom border
- Company/school + location on one line (right-tab at 7.19")
- Title/degree + dates on next line (right-tab)
- Proper bullet list indentation (left=270, hanging=280 twips)

Layout order: EDUCATION -> EXPERIENCE -> TECHNICAL SKILLS -> ADDITIONAL INFORMATION
"""

from __future__ import annotations

from pathlib import Path  # noqa: TC003
from typing import Any

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

from pipelines.job_agent.models.resume_tailoring import (
    TailoredResumeDocument,  # noqa: TC001
)

logger = structlog.get_logger(__name__)

_FONT_NAME = "Times New Roman"
_FONT_SIZE_HALF_PTS = 23  # 11.5pt in half-point units
_FONT_SIZE_PT = Pt(11.5)
_LINE_SPACING = Twips(200)  # 10pt atLeast
_RIGHT_TAB_POS = Inches(7.19)
_BORDER_SZ = "12"  # 1.5pt black bottom border for section headers
_BULLET_LEFT = 270  # twips
_BULLET_HANGING = 280  # twips


def render_resume_docx(doc: TailoredResumeDocument, output_path: Path) -> Path:
    """Build a .docx resume from *doc* and write it to *output_path*.

    Section order matches the Kokomoor template:
    EDUCATION -> EXPERIENCE -> TECHNICAL SKILLS -> ADDITIONAL INFORMATION
    """
    document = Document()
    _setup_page(document)
    _set_default_font(document)
    _ensure_numbering(document)

    _render_header(document, doc)
    _add_spacer(document)

    if doc.education:
        _render_education(document, doc)

    _add_spacer(document)

    if doc.experience:
        _render_experience(document, doc)

    _add_spacer(document)

    if doc.skills_highlight:
        _render_skills(document, doc)

    _add_spacer(document)

    if doc.additional_info or doc.clearance or doc.supplementary_projects:
        _render_additional_info(document, doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    logger.info("resume_rendered", path=str(output_path))

    preview_path = output_path.with_suffix(".md")
    _render_markdown_preview(doc, preview_path)

    return output_path


# -- page setup -----------------------------------------------------------


def _setup_page(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def _ensure_numbering(document: Any) -> None:
    """Inject a bullet-list numbering definition if the document lacks one.

    Defines abstractNumId 0 with a bullet character at indent level 0,
    then binds numId 1 to it. This makes ``numId=1`` available for all
    bullet paragraphs created by ``_set_bullet_indent``.
    """
    numbering_part = document.part.numbering_part
    numbering_elem = numbering_part.numbering_definitions._numbering

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "0")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Symbol")
    r_fonts.set(qn("w:hAnsi"), "Symbol")
    r_fonts.set(qn("w:hint"), "default")
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract_num.append(lvl)
    numbering_elem.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "1")
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), "0")
    num.append(abstract_ref)
    numbering_elem.append(num)


def _set_default_font(document: Any) -> None:
    """Set document-level default to Times New Roman."""
    rpr_default = document.styles.element.find(qn("w:docDefaults"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:docDefaults")
        document.styles.element.insert(0, rpr_default)
    rpr_def = rpr_default.find(qn("w:rPrDefault"))
    if rpr_def is None:
        rpr_def = OxmlElement("w:rPrDefault")
        rpr_default.append(rpr_def)
    rpr = rpr_def.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_def.append(rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), _FONT_NAME)
    fonts.set(qn("w:hAnsi"), _FONT_NAME)
    fonts.set(qn("w:eastAsia"), _FONT_NAME)
    fonts.set(qn("w:cs"), _FONT_NAME)


# -- paragraph builders ---------------------------------------------------


def _make_para(document: Any) -> Any:
    """Create a paragraph with standard line spacing and zero space before/after."""
    para = document.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = _LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    return para


def _add_run(
    para: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    caps: bool = False,
    color_rgb: str | None = None,
) -> Any:
    """Add a run with standard font size and optional formatting."""
    run = para.add_run(text)
    run.font.size = _FONT_SIZE_PT
    run.font.name = _FONT_NAME
    run.bold = bold
    run.italic = italic
    if caps:
        run.font.all_caps = True
    if color_rgb:
        rpr = run._element.get_or_add_rPr()
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), color_rgb)
        rpr.append(color_elem)
    return run


def _add_right_tab(para: Any) -> None:
    """Add a right-aligned tab stop at 7.19 inches."""
    pf = para.paragraph_format
    pf.tab_stops.add_tab_stop(_RIGHT_TAB_POS, WD_TAB_ALIGNMENT.RIGHT)


def _add_bottom_border(para: Any, color: str = "000000", sz: str = _BORDER_SZ) -> None:
    """Add a bottom border to a paragraph."""
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_spacer(document: Any) -> None:
    """Add a non-breaking space paragraph as a visual separator."""
    para = _make_para(document)
    _add_run(para, "\u00a0")


def _add_bullet_para(document: Any, text: str) -> None:
    """Add a bulleted paragraph matching the reference template indentation."""
    para = _make_para(document)
    _set_bullet_indent(para)
    _add_run(para, text)


def _set_bullet_indent(para: Any) -> None:
    """Apply bullet-style indentation and a bullet character prefix via numbering XML."""
    p_pr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(_BULLET_LEFT))
    ind.set(qn("w:hanging"), str(_BULLET_HANGING))
    p_pr.append(ind)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)


# -- section renderers -----------------------------------------------------


def _render_header(document: Any, doc: TailoredResumeDocument) -> None:
    """Render name and contact info, centered with invisible bottom border."""
    name_para = _make_para(document)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(name_para, color="FFFFFF", sz="6")
    _add_run(name_para, doc.name, bold=True, caps=True)

    contact_para = _make_para(document)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(contact_para, color="FFFFFF", sz="6")

    parts = [p for p in [doc.location, doc.phone, doc.email, doc.linkedin] if p]
    for i, part in enumerate(parts):
        if i > 0:
            _add_run(contact_para, "\u00a0|\u00a0", color_rgb="000000")
        is_url = part.startswith("http") or "linkedin.com" in part
        _add_run(contact_para, part, color_rgb="0000FF" if is_url else None)


def _render_section_header(document: Any, title: str) -> None:
    """Render a section header: bold, ALL CAPS, black bottom border."""
    para = _make_para(document)
    _add_bottom_border(para)
    _add_run(para, title, bold=True, caps=True)


def _render_education(document: Any, doc: TailoredResumeDocument) -> None:
    _render_section_header(document, "education")

    for edu in doc.education:
        school_para = _make_para(document)
        _add_right_tab(school_para)
        _add_run(school_para, edu.school, bold=True, caps=True)
        if edu.location:
            _add_run(school_para, "\t")
            _add_run(school_para, edu.location)

        degree_para = _make_para(document)
        _add_right_tab(degree_para)
        _add_run(degree_para, edu.degree, italic=True)
        if edu.graduation:
            _add_run(degree_para, "\t")
            _add_run(degree_para, edu.graduation)

        for bullet in edu.bullets:
            _add_bullet_para(document, bullet.text)


def _render_experience(document: Any, doc: TailoredResumeDocument) -> None:
    _render_section_header(document, "experience")

    for idx, exp in enumerate(doc.experience):
        if idx > 0:
            _add_spacer(document)

        company_para = _make_para(document)
        _add_right_tab(company_para)
        _add_run(company_para, exp.company, bold=True, caps=True)
        if exp.location:
            _add_run(company_para, "\t")
            _add_run(company_para, exp.location)

        if exp.subtitle:
            sub_para = _make_para(document)
            _add_run(sub_para, exp.subtitle, italic=True)

        title_para = _make_para(document)
        _add_right_tab(title_para)
        _add_run(title_para, exp.title, bold=True, italic=True)
        if exp.dates:
            _add_run(title_para, "\t")
            _add_run(title_para, exp.dates)

        for bullet in exp.bullets:
            _add_bullet_para(document, bullet.text)


def _render_skills(document: Any, doc: TailoredResumeDocument) -> None:
    _render_section_header(document, "technical skills")

    skills_text = ", ".join(doc.skills_highlight)
    _add_bullet_para(document, skills_text)


def _render_additional_info(document: Any, doc: TailoredResumeDocument) -> None:
    _render_section_header(document, "additional information")

    for item in doc.additional_info:
        _add_bullet_para(document, item)

    # Supplementary projects are rendered as one bulleted line per
    # project: "Name — descriptive text (url)". Personal projects like
    # kokomoor-platform live here rather than in Experience so they do
    # not compete with real work history for section real estate.
    for proj in doc.supplementary_projects:
        line = proj.text
        if proj.url and proj.url not in line:
            line = f"{line} ({proj.url})"
        _add_bullet_para(document, line)


# -- markdown preview ------------------------------------------------------


def _render_markdown_preview(doc: TailoredResumeDocument, path: Path) -> None:
    """Write a Markdown preview of the tailored resume alongside the .docx."""
    lines: list[str] = []

    contact = " | ".join(p for p in [doc.location, doc.phone, doc.email, doc.linkedin] if p)
    lines.append(f"# {doc.name}")
    if contact:
        lines.append(f"_{contact}_")
    lines.append("")

    if doc.education:
        lines.append("## EDUCATION")
        lines.append("")
        for edu in doc.education:
            header = f"**{edu.school}**"
            if edu.location:
                header += f" — {edu.location}"
            lines.append(header)
            degree_line = f"_{edu.degree}_"
            if edu.graduation:
                degree_line += f" ({edu.graduation})"
            lines.append(degree_line)
            for bullet in edu.bullets:
                lines.append(f"- {bullet.text}")
            lines.append("")

    if doc.experience:
        lines.append("## EXPERIENCE")
        lines.append("")
        for exp in doc.experience:
            header = f"**{exp.company}**"
            if exp.location:
                header += f" — {exp.location}"
            lines.append(header)
            title_line = f"**_{exp.title}_**"
            if exp.dates:
                title_line += f" | {exp.dates}"
            lines.append(title_line)
            for bullet in exp.bullets:
                lines.append(f"- {bullet.text}")
            lines.append("")

    if doc.skills_highlight:
        lines.append("## TECHNICAL SKILLS")
        lines.append("")
        lines.append(f"- {', '.join(doc.skills_highlight)}")
        lines.append("")

    if doc.additional_info or doc.clearance or doc.supplementary_projects:
        lines.append("## ADDITIONAL INFORMATION")
        lines.append("")
        for item in doc.additional_info:
            lines.append(f"- {item}")
        for proj in doc.supplementary_projects:
            line = proj.text
            if proj.url and proj.url not in line:
                line = f"{line} ({proj.url})"
            lines.append(f"- {line}")
        lines.append("")

    try:
        path.write_text("\n".join(lines), encoding="utf-8")
        logger.info("resume_preview_rendered", path=str(path))
    except Exception:
        logger.warning("resume_preview_failed", path=str(path), exc_info=True)
